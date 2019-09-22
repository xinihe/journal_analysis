# -*- coding: utf-8 -*-
"""
Created on Thu Feb 28 15:23:04 2019

@author: Ni He
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from wordcloud import WordCloud
from sqlalchemy import create_engine
import ChineseSurname as cs
from docx import Document
from docx.shared import Cm,Pt,RGBColor
import os
import mysql.connector

def gen_wordcloud(text):
    wordcloud = WordCloud(background_color="white", #设置背景为白色，默认为黑色
                      width=1500, #设置图片的宽度 
                      height=960, #设置图片的高度 
                      margin=10 #设置图片的边缘 
                      )
    wordcloud.generate(text)
    return wordcloud

def con_mysql(sql):  
    conn = mysql.connector.connect(host="*****",port=3306,user="root",\
                           password="*****",database="cnki",charset="utf8")
    cur = conn.cursor()
    cur.execute(sql)
    info = cur.fetchall()
    cur.close
    return info

def get_info(table_name):

    engine = create_engine('mysql+pymysql://root:******@10.23.0.2:3306/fintech?charset=utf8')
    content_list = 'pkeywords, ppage, volume, pdate, authors, paddress, title' #where_cond = "pdate like '%2018'" 
    sql_cmd = "select {} from {} ".format(content_list, table_name)
    # retrieval data from database
    try:
        df = pd.read_sql(sql=sql_cmd, con=engine)
    except:
        if input('The journal may not be in collection. Do you want to retry?[y or n]') == 'y':
            get_info()
    # Validation of retrievaled data
    invalid_rows = []
    for l in range(df.shape[0]):
        for s in range(df.shape[1]):
            if s in [0,1,2,3,4,5] and df.iloc[l,s] == 'Null':
                invalid_rows.append(l)
                print('Line %s has been deleted due to NULL.'%l)
                break
            #if s == 1 and ( '' in df.iloc[l,s].strip().split('-') or not len(df.iloc[l,s].strip().split('-')) == 2 ):
            if s == 1 and not all([f.isdigit() for f in df.iloc[l,s].strip().split('-')]): # if each element is digit
                invalid_rows.append(l)
                print('Line %s has been deleted due to inappropriate pagenumber.'%l)
                break
#            if s == 2 and not df.iloc[l,s].isdigit():
#                invalid_rows.append(l)
#                print('Line %s has been deleted due to inappropriate volume.'%l)
#                break
            if s == 3 and not df.iloc[l,s].strip().split(' ')[-1].isdigit():  # the last digit is a valid year
                invalid_rows.append(l)
                print('Line %s has been deleted due to inappropriate date.'%l)
                break
    return df.drop(invalid_rows), len(invalid_rows)

def add_info(df):
    df['pages'] = [int(ppages.split('-')[1].strip()) - int(ppages.split('-')[0].strip()) + 1 for ppages in df['ppage']]
    df['year'] = [pdate.split(' ')[-1].replace('\n','') for pdate in df['pdate']]
    df['month'] = [pdate.split(' ')[0].replace('\n','') for pdate in df['pdate']]

    df['institute'] = [addr.split(',')[0].strip() for addr in df['paddress']]
    df['author1st'] = [au.split(',')[0] for au in df['authors']]
    
    country_list = [addr.split(',')[-1].strip().replace('.','') for addr in df['paddress']]
    for i in range(len(country_list)):
        if 'USA' in country_list[i]:
            country_list[i] = 'USA'
        if 'China' in country_list[i]:
            country_list[i] = 'China'
    df['country'] = country_list
    return df

def create_statistics(df):
    info = []
    years = list(set(df['year']))
    years.sort()
    num_of_keywords = 10
    for y in years:
        dic = {}
        dic['year'] = y
        dic['num_pubs'] = list(df['year']).count(y)
        dic['aver_page'] = round(np.mean(list(df[df['year'] == y]['pages'])))
        kk = ";".join(df[df['year'] == y]['pkeywords']).replace('KeyWords Plus:','').replace('Null','').lower().split(';')
        dic['keywords'] = [k.strip() for k in kk]
        dic['percent_cn_univ'] = round(100 * len(df[(df['year'] == y) & (df['country'] == 'China')]['country'])/dic['num_pubs'])
        author_list = list(df[df['year'] == y]['author1st'])
        cn_names = [cs.findsurname(au) for au in author_list if cs.findsurname(au) is not 'NA']
        dic['percent_cn_author'] = round(100 * len(cn_names)/dic['num_pubs'])
        dic['list_keyw'] = word_counter_in_list(word_counter(dic['keywords'],num_of_keywords))
        
        info.append(dic)
    return info

def word_counter(txt_list,num_of_word = 0):
    word = []
    counter = {}   
    for txt in txt_list:
        txt = txt.strip()
        if len(txt) == 0:
            continue
        if not txt in word:
            word.append(txt)
        if not txt in counter.keys():
            counter[txt] = 1
        else:
            counter[txt] += 1
    
    counter_list = sorted(counter.items(), key = lambda x: x[1], reverse = True)  # sort by the value of dic
    
    if num_of_word == 0:
        return counter_list
    else:
        return counter_list[:num_of_word]

def word_counter_in_list(counter_list):
    return [' - '.join([str(x) for x in list(k)]) for k in counter_list]

def plot_wordcloud(info,table_name,cwd):
    plt.figure()
    info.sort(key=lambda x:x['year'])
    for i in range(len(info)):
        text = ' '.join(info[i]['keywords'])
        wc = gen_wordcloud(text)
        px = plt.subplot(np.ceil(np.sqrt(len(info))),np.floor(np.sqrt(len(info))),i+1)    

        px.imshow(wc)
        px.set_title(info[i]['year'])
        px.axis("off")

    #plt.show()
    plt.savefig(cwd+'\\'+table_name+'_wc.jpg', format ='jpg')

def plot_country_distribution(df,table_name, cwd):
    num_of_countries = 20
    plt.figure()
    countries = df.country.tolist()
    coun_hist = word_counter(countries)
    label_list = [c[0] for c in coun_hist]
    #label_list[label_list.index('Peoples R China')] = 'China'   # name is too long
    num_list = [c[1] for c in coun_hist]
    # plot image
    plt.bar(range(len(num_list[:num_of_countries])), num_list[:num_of_countries])
    index = list(range(len(num_list[:num_of_countries])))
    plt.xticks(index, label_list[:num_of_countries],rotation=80)
    
    plt.title('Number of Publications by Different Countries')
    #plt.show()
    plt.savefig(cwd+'\\'+table_name+'_cd.jpg', format ='jpg', bbox_inches = 'tight')
    return label_list, num_list
    
def plot_institute_distribution(df,table_name, cwd):
    num_of_institute = 20
    plt.figure()
    institutes = df.institute.tolist()
    ins_hist = word_counter(institutes)
    label_list = [c[0] for c in ins_hist]
    num_list = [c[1] for c in ins_hist]
    
    plt.bar(range(len(num_list[:num_of_institute])), num_list[:num_of_institute])
    index = list(range(len(num_list[:num_of_institute])))
    plt.xticks(index, label_list[:num_of_institute],rotation=80)
    plt.title('Number of Publications by different universities')
    #plt.show()
    plt.savefig(cwd+'\\'+table_name+'_id.jpg', format ='jpg', bbox_inches = 'tight')
    return label_list, num_list

def cal_cn_univ_distribution(df):
    institutes = df.institute.tolist()
    countries = df.country.tolist()
    cn_inst_index = [i for i,x in enumerate(countries) if 'China' in x]
    cn_inst = [x for i,x in enumerate(institutes) if i in cn_inst_index]
    return word_counter(cn_inst)
    
        
def any_univ(df, univ = 'Zhejiang Gongshang'):
    institutes = df.institute.tolist()
    index = [i for i in range(len(institutes)) if univ in institutes[i]]
    return index

def general_info(journal_name):
    #like_name = '%' + ('%').join(journal_name.replace('&','').replace('_',' ').replace('and', '').split(' ')) + '%'
    like_name = (' ').join(journal_name.replace('_',' ').split(' '))


    sql_cmd = "select * from ibs_toplist where JournalName = '{}'".format(like_name)
    sql_ssci = "select FullJournalTitle, Quartile from ssci_details where FullJournalTitle = '{}'".format(like_name)
    sql_sci = "select FullJournalTitle, Quartile from scie_details where FullJournalTitle = '{}'".format(like_name)
    sql_ssci_details = "select * from journal_list_ssci where title = '{}'".format(like_name)
    sql_sci_details = "select * from journal_list_scie where title = '{}'".format(like_name)
    
    info_ibs = con_mysql(sql_cmd)
    info_sci = con_mysql(sql_sci)
    info_ssci = con_mysql(sql_ssci)
    info_sci_details = con_mysql(sql_sci_details)
    info_ssci_details = con_mysql(sql_ssci_details)
    
    journal_info = {}
    journal_info['JournalName'] = journal_name
    if info_ibs:
        journal_info['UnivRank'] = info_ibs[0][4]
        journal_info['IBSRank'] = info_ibs[0][3]
    else:
        journal_info['UnivRank'] = 'No'
        journal_info['IBSRank'] = 'No'
    if info_sci:
        journal_info['sci'] = info_sci[0][1]
    else:
        journal_info['sci'] = 'No'
    if info_ssci:
        journal_info['ssci'] = info_ssci[0][1]
    else:
        journal_info['ssci'] = 'No'
    
    if info_sci_details:
        journal_info['issn'] =    info_sci_details[0][2]
        journal_info['freq'] =    info_sci_details[0][3]
        journal_info['press'] =   info_sci_details[0][4]
        journal_info['city'] =    info_sci_details[0][5]
        journal_info['country'] = info_sci_details[0][6]
        journal_info['address'] = info_sci_details[0][7]
    elif info_ssci_details:
        journal_info['issn'] =    info_ssci_details[0][2]
        journal_info['freq'] =    info_ssci_details[0][3]
        journal_info['press'] =   info_ssci_details[0][4]
        journal_info['city'] =    info_ssci_details[0][5]
        journal_info['country'] = info_ssci_details[0][6]
        journal_info['address'] = info_ssci_details[0][7]
    else:
        journal_info['issn'] = 'NA'
        journal_info['freq'] = 'NA'
        journal_info['press'] = 'NA'
        journal_info['city'] = 'NA'
        journal_info['country'] = 'NA'
        journal_info['address'] = 'NA'
    return journal_info
            

def generate_word_report(journal_name, table_name,info,df,cwd,invalid_lines):   
    document = Document()
    document.add_heading(journal_name, 0)
    
    document.add_heading('Journal Analysis...', level=1) 
    journal_info = general_info(journal_name)
    document.add_paragraph('SCI: '+journal_info['sci'])
    document.add_paragraph('SSCI: '+journal_info['ssci'])
    document.add_paragraph('University Top Journal: '+journal_info['UnivRank'])
    document.add_paragraph('International Business School Top Journal: '+journal_info['IBSRank'])
    document.add_paragraph('ISSN: '+journal_info['issn'])
    document.add_paragraph('Publication Frequency: '+journal_info['freq'].upper())
    document.add_paragraph('Press: '+journal_info['press'].upper())
    document.add_paragraph('City: '+journal_info['city'].upper())
    document.add_paragraph('Country: '+journal_info['country'].upper())
    document.add_paragraph('Address: '+journal_info['address'].upper())
    

    
    #document.add_paragraph('Journal Statistics')
    document.add_heading('Journal Statistics: ({} out of {} entries have been deleted)'.format(invalid_lines, invalid_lines+len(df)), level=1) 
    
    # add table ------------------
    rows = 1
    cols = 5
    table = document.add_table(rows=rows, cols=cols, style = "Light List Accent 1") 
    # populate header row -----
    column_names = ['Year','AveragePage','NumPublication','ChineseAuthor','ChineseUniv']
    heading_cells = table.rows[0].cells
    for i in range(len(column_names)):
        heading_cells[i].text = column_names[i]
    
    for item in info:
        cells = table.add_row().cells
        cells[0].text = item['year']
        cells[1].text = str(int(item['aver_page']))
        cells[2].text = str(item['num_pubs'])
        cells[3].text = str(item['percent_cn_author'])+'%'
        cells[4].text = str(item['percent_cn_univ'])+'%'

    document.add_heading('Journal Keywords Distributions - Top 10 most popular Keywords', level=1) 
    # add table ------------------
    num_of_keywords = 10
    rows = num_of_keywords + 1
    cols = len(info)
    table = document.add_table(rows=rows, cols=cols,style = "Table Grid") 
    # populate header row -----
    heading_cells = table.rows[0].cells
    for i in range(len(info)):
        heading_cells[i].text = info[i]['year']
    
    for c in range(cols):
        for r in range(1,rows): # The first row has been populated, start from the second row
            run=table.cell(r,c).paragraphs[0].add_run(info[c]['list_keyw'][r-1])
            run.font.size=Pt(8)
    #        cell = table.cell(r, c)
    #        cell.text = info[c]['list_keyw'][r-1]
    
    
    
    # add images ------------------
    document.add_heading('Journal Keywords Cloud by Years', level=1) 
    plot_wordcloud(info,table_name,cwd)
    document.add_picture(cwd+'\\'+table_name+'_wc.jpg', width=Cm(17))
    
    document.add_page_break() 
    
    
    # add a new page ------------------
    document.add_heading('Country Distribution', level=1) 
    labels, nums = plot_country_distribution(df,table_name, cwd)
    document.add_picture(cwd+'\\'+table_name+'_cd.jpg', width=Cm(15))
    
    document.add_heading('List of most productive country...', level=2) 
    for i in range(10):
        p = document.add_paragraph('Productive country ', style='ListBullet')
        p.add_run('No. '+str(i+1)+' is ')
        p.add_run(labels[i]).bold = True
        p.add_run(', which has ')
        p.add_run(str(nums[i])).bold = True
        p.add_run(' publications.')
    
    if 'China' not in labels:
        document.add_heading('There is NOT publication from any Chinese university...', level=2) 
    else:
        i = labels.index('China')
        document.add_heading('China is ranked at No. '+str(i+1)+' and has '+str(nums[i])+' Publications in this journal.', level=2) 
        
    document.add_page_break() 
    
    # add a new page ------------------
    document.add_heading('Institute Distribution', level=1) 
    labels, nums = plot_institute_distribution(df,table_name, cwd)
    document.add_picture(cwd+'\\'+table_name+'_id.jpg', width=Cm(15))
    
    document.add_heading('List of most productive institute...', level=2) 
    for i in range(10):
        p = document.add_paragraph('Productive institute ', style='ListBullet')
        p.add_run('No. '+str(i+1)+' is ')
        p.add_run(labels[i]).bold = True
        p.add_run(', which has ')
        p.add_run(str(nums[i])).bold = True
        p.add_run(' publications.')
    
    # Top chinese universities ------
    cn_ins_hist = cal_cn_univ_distribution(df)
    if not cn_ins_hist:
        document.add_heading('We are on the way...', level=2) 
    else:
        document.add_heading('List of most productive Chinese institute...', level=2) 
        label_list = [c[0] for c in cn_ins_hist]
        num_list = [c[1] for c in cn_ins_hist]
        for i in range(min(10,len(cn_ins_hist))): # display top 10 (if available) universities
            p = document.add_paragraph('Productive institute ', style='ListBullet')
            p.add_run('No. '+str(i+1)+' is ')
            p.add_run(label_list[i]).bold = True
            p.add_run(', which has ')
            p.add_run(str(num_list[i])).bold = True
            p.add_run(' publications.')
    
 
    # add a new page ------------------
    # add table ------------------
    univs = ['Zhejiang Gongshang Univ','Zhejiang Univ Finance & Econ','Shanghai Univ Finance & Econ',
             'Jiangxi Univ Finance & Econ', 'Cent Univ Finance & Econ', 'Dongbei Univ Finance & Econ',
             'Southwestern Univ Finance & Econ','Univ Int Business & Econ','Zhongnan Univ Econ & Law']
    for univ in univs:
        document.add_heading('Publication Record of '+univ, level=1) 
        index = any_univ(df, univ)
        if not index:
            document.add_paragraph(univ+' has not been contributed to this journal yet...')
          
        else:  
            rows = 1
            cols = 5
            table = document.add_table(rows=rows, cols=cols,style = "Light List Accent 1") 
            #table.autofit=True
            table.cell(0,0).width=Cm(0.5)
            table.cell(0,1).width=Cm(1)
            table.cell(0,2).width=Cm(2.5)
            table.cell(0,3).width=Cm(2.5)
            table.cell(0,4).width=Cm(8)
            # populate header row -----
            column_names = ['ID','Year','Author(s)','School','Title']
            heading_cells = table.rows[0].cells
            for i in range(len(column_names)):
                heading_cells[i].text = column_names[i]
            
            i = 0
            for it in range(len(index)):
                item = df.iloc[index[it],:]
                i += 1
                cells = table.add_row().cells
                cells[0].text = str(i)
                cells[1].text = item['year']
                cells[2].text = item['authors']
                cells[3].text = item['paddress'].split(',')[1].strip()
                cells[4].text = item['title']
        
        
    # save to doc  ------------------
    document.save(cwd+'\\'+table_name+'.docx')
         


def main_analysis_func(journal_name):
    #journal_name = input('Please input the full name of the journal(i.e. Review of Financial Studies):') 
    table_name = '_'.join(journal_name.lower().strip().split(' '))
    cwd = os.getcwd()+'\\'+table_name
    if not os.path.exists(cwd):
        os.makedirs(cwd)
    
    df,invalid_lines = get_info(table_name)
    df = add_info(df)
    info = create_statistics(df)
    generate_word_report(journal_name, table_name, info, df, cwd, invalid_lines)
    
    
if __name__ == '__main__':
    journal_name = input('Please input the full name of the journal(i.e. Review of Financial Studies):') 
    main_analysis_func(journal_name)






#data20163 = pd.read_csv('2016.csv', index_col = 0, skiprows = 1, header = 0, encoding='gbk',\
#            names = ['Dept', 'Title', 'Author', 'Pos', 'OtherAuthor', 'Index', 'Date', 'Rank','Journal', 'PubDate', 'PubYear','Note'])
